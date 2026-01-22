"""Export photos to Word document"""

import io
import math
from typing import List, Callable, Optional

from PIL import Image
from PyQt5.QtCore import QThread, pyqtSignal
from docx import Document
from docx.shared import Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH

from ..models import PhotoItem
from ..config import WordExportConfig

class WordExporter(QThread):
    """Thread for Word export"""

    progress = pyqtSignal(int)
    finished = pyqtSignal(str)
    error = pyqtSignal(str)

    def __init__(self, photos: List[PhotoItem], path: str, photos_per_page: int, image_size: str = 'full'):
        super().__init__()
        self.photos = photos
        self.path = path
        self.ppp = photos_per_page
        self.image_size = image_size  # 'half', 'three_quarter', or 'full'
        self.config = WordExportConfig()

    def run(self) -> None:
        """Execute the export"""
        try:
            self._generate_word()
            self.finished.emit(self.path)
        except Exception as e:
            self.error.emit(str(e))

    def _generate_word(self) -> None:
        """Generate the Word document"""
        doc = Document()

        # Configure page margins
        page_margin = self.config.PAGE_MARGIN_MM
        for section in doc.sections:
            section.top_margin = Mm(page_margin)
            section.bottom_margin = Mm(page_margin)
            section.left_margin = Mm(page_margin)
            section.right_margin = Mm(page_margin)

        # Layout based on photos per page
        cols, rows = self.config.LAYOUTS.get(self.ppp, (2, 3))

        # Available area (A4 = 210x297mm)
        available_w_mm = 210 - (page_margin * 2)
        available_h_mm = 297 - (page_margin * 2)

        # Gap between photos (FIXED, same horizontal and vertical)
        gap_mm = self.config.GAP_MM

        # Calculate cell width (fixed)
        cell_w_mm = (available_w_mm - gap_mm * (cols - 1)) / cols

        # Apply size factor to cells only (not to gaps)
        size_factor = self.config.IMAGE_SIZES.get(self.image_size, 1.0)
        cell_w_mm = cell_w_mm * size_factor

        # Convert mm to pixels
        mm_to_px = self.config.DPI / 25.4
        cell_w_px = int(cell_w_mm * mm_to_px)
        gap_px = int(gap_mm * mm_to_px)

        total = len(self.photos)
        num_pages = math.ceil(total / self.ppp)

        for page_idx in range(num_pages):
            if page_idx > 0:
                doc.add_page_break()

            start = page_idx * self.ppp
            end = min(start + self.ppp, total)
            page_photos = self.photos[start:end]

            # Calculate row heights based on tallest photo in each row
            row_heights_px = []
            for row_idx in range(rows):
                row_start = row_idx * cols
                row_end = min(row_start + cols, len(page_photos))
                row_photos = page_photos[row_start:row_end]

                if not row_photos:
                    continue

                # Find tallest photo in this row
                max_h_px = 0
                for photo in row_photos:
                    try:
                        with Image.open(photo.path) as img:
                            img_w, img_h = img.size
                            img_ratio = img_w / img_h
                            new_w = cell_w_px
                            new_h = int(new_w / img_ratio)
                            max_h_px = max(max_h_px, new_h)
                    except Exception as e:
                        print(f"Error reading photo {photo.path}: {e}")

                row_heights_px.append(max_h_px)

            # Total composite height
            composite_h_px = sum(row_heights_px) + gap_px * (len(row_heights_px) - 1)
            composite_w_px = cols * cell_w_px + (cols - 1) * gap_px

            # Create composite image (white background)
            composite = Image.new('RGB', (composite_w_px, composite_h_px), (255, 255, 255))

            y_pos = 0
            for row_idx in range(rows):
                row_start = row_idx * cols
                row_end = min(row_start + cols, len(page_photos))
                row_photos = page_photos[row_start:row_end]

                if not row_photos:
                    continue

                row_h_px = row_heights_px[row_idx]

                for col_idx in range(len(row_photos)):
                    photo = row_photos[col_idx]
                    x = col_idx * (cell_w_px + gap_px)

                    # Place photo
                    self._place_photo(composite, photo, x, y_pos, cell_w_px, row_h_px)

                    # Update progress
                    self.progress.emit(int((start + row_start + col_idx + 1) / total * 100))

                y_pos += row_h_px + gap_px

            # Insert composite image into document
            self._insert_composite(doc, composite, composite_w_px / mm_to_px, composite_h_px / mm_to_px)

        doc.save(self.path)

    def _place_photo(
        self,
        composite: Image.Image,
        photo: PhotoItem,
        x: int,
        y: int,
        cell_w: int,
        cell_h: int
    ) -> None:
        """Place a photo in the composite image (FIT width, height auto, centered vertically)"""
        try:
            with Image.open(photo.path) as img:
                # Apply rotation
                if photo.rotation:
                    img = img.rotate(-photo.rotation, expand=True)

                # Convert to RGB
                if img.mode != 'RGB':
                    if img.mode in ('RGBA', 'LA', 'P'):
                        bg = Image.new('RGB', img.size, (255, 255, 255))
                        if img.mode == 'P':
                            img = img.convert('RGBA')
                        if 'A' in img.getbands():
                            bg.paste(img, mask=img.split()[-1])
                        else:
                            bg.paste(img)
                        img = bg
                    else:
                        img = img.convert('RGB')

                # FIT width, height auto
                img_w, img_h = img.size
                img_ratio = img_w / img_h
                new_w = cell_w
                new_h = int(new_w / img_ratio)

                # Resize
                resample = Image.LANCZOS if hasattr(Image, 'LANCZOS') else Image.ANTIALIAS
                img_resized = img.resize((new_w, new_h), resample)

                # Center vertically in cell
                paste_x = x
                paste_y = y + (cell_h - new_h) // 2

                composite.paste(img_resized, (paste_x, paste_y))

        except Exception as e:
            print(f"Error placing photo {photo.path}: {e}")

    def _insert_composite(
        self,
        doc: Document,
        composite: Image.Image,
        width_mm: float,
        height_mm: float
    ) -> None:
        """Insert the composite image into the document"""
        buf = io.BytesIO()
        composite.save(buf, format='JPEG', quality=self.config.JPEG_QUALITY)
        buf.seek(0)

        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_before = Mm(0)
        para.paragraph_format.space_after = Mm(0)
        para.add_run().add_picture(buf, width=Mm(width_mm), height=Mm(height_mm))

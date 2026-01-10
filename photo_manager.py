#!/usr/bin/env python3
"""
Photo Manager - Application de gestion de photos pour generer des documents Word
Compatible Windows et macOS (PyQt5)
"""

import sys
import os
import math
import io
from typing import List, Optional
from dataclasses import dataclass, field

from PyQt5.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
    QLabel, QPushButton, QRadioButton, QButtonGroup, QScrollArea,
    QGridLayout, QFileDialog, QMessageBox, QProgressDialog, QFrame,
    QSizePolicy, QDialog
)
from PyQt5.QtCore import Qt, QThread, pyqtSignal, QSize
from PyQt5.QtGui import QPixmap, QImage, QFont

from PIL import Image
from docx import Document
from docx.shared import Mm
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SUPPORTED_FORMATS = ('.jpg', '.jpeg', '.png', '.JPG', '.JPEG', '.PNG')
THUMB_SIZE = (100, 100)
PHOTOS_PER_VIEW = 50


@dataclass
class PhotoItem:
    """Photo avec metadonnees"""
    path: str
    rotation: int = 0
    _pixmap: Optional[QPixmap] = field(default=None, repr=False)

    def rotate(self):
        self.rotation = (self.rotation + 90) % 360
        self._pixmap = None

    def get_pixmap(self) -> Optional[QPixmap]:
        if self._pixmap:
            return self._pixmap
        try:
            with Image.open(self.path) as img:
                if self.rotation:
                    img = img.rotate(-self.rotation, expand=True)
                if img.mode != 'RGB':
                    img = img.convert('RGB')
                img.thumbnail(THUMB_SIZE, Image.LANCZOS if hasattr(Image, 'LANCZOS') else Image.ANTIALIAS)

                # Convert PIL to QPixmap
                data = img.tobytes("raw", "RGB")
                qimg = QImage(data, img.width, img.height, 3 * img.width, QImage.Format_RGB888)
                self._pixmap = QPixmap.fromImage(qimg)
                return self._pixmap
        except Exception as e:
            print(f"Error loading {self.path}: {e}")
            return None

    def clear(self):
        self._pixmap = None


class ImageViewerDialog(QDialog):
    """Modal pour afficher une photo en grand"""

    def __init__(self, photo: PhotoItem, parent=None):
        super().__init__(parent)
        self.photo = photo
        self.setWindowTitle(os.path.basename(photo.path))
        self.setModal(True)

        # Taille max = 90% de l'ecran
        screen = QApplication.primaryScreen().geometry()
        max_w = int(screen.width() * 0.9)
        max_h = int(screen.height() * 0.9)

        layout = QVBoxLayout(self)
        layout.setContentsMargins(10, 10, 10, 10)

        # Image label
        self.img_label = QLabel()
        self.img_label.setAlignment(Qt.AlignCenter)
        layout.addWidget(self.img_label)

        # Bouton fermer
        close_btn = QPushButton("Fermer")
        close_btn.clicked.connect(self.close)
        layout.addWidget(close_btn)

        # Charger l'image en grand
        self._load_full_image(max_w, max_h - 50)

    def _load_full_image(self, max_w, max_h):
        try:
            with Image.open(self.photo.path) as img:
                if self.photo.rotation:
                    img = img.rotate(-self.photo.rotation, expand=True)
                if img.mode != 'RGB':
                    img = img.convert('RGB')

                # Redimensionner pour tenir dans la fenetre
                img_w, img_h = img.size
                scale = min(max_w / img_w, max_h / img_h, 1.0)
                new_w = int(img_w * scale)
                new_h = int(img_h * scale)

                resample = Image.LANCZOS if hasattr(Image, 'LANCZOS') else Image.ANTIALIAS
                img = img.resize((new_w, new_h), resample)

                # Convertir en QPixmap
                data = img.tobytes("raw", "RGB")
                qimg = QImage(data, img.width, img.height, 3 * img.width, QImage.Format_RGB888)
                pixmap = QPixmap.fromImage(qimg)

                self.img_label.setPixmap(pixmap)
                self.resize(new_w + 20, new_h + 70)
        except Exception as e:
            self.img_label.setText(f"Erreur: {e}")


class PhotoCard(QFrame):
    """Carte photo compacte"""

    def __init__(self, photo: PhotoItem, index: int, on_delete, on_rotate):
        super().__init__()
        self.photo = photo
        self.index = index
        self.on_delete = on_delete
        self.on_rotate = on_rotate

        self.setFrameStyle(QFrame.Box | QFrame.Raised)
        self.setFixedSize(130, 160)
        self.setCursor(Qt.PointingHandCursor)

        layout = QVBoxLayout(self)
        layout.setContentsMargins(5, 5, 5, 5)
        layout.setSpacing(2)

        # Image label (cliquable)
        self.img_label = QLabel()
        self.img_label.setFixedSize(100, 100)
        self.img_label.setAlignment(Qt.AlignCenter)
        self.img_label.setStyleSheet("background-color: #f0f0f0;")
        self.img_label.setCursor(Qt.PointingHandCursor)
        self.img_label.mousePressEvent = self._show_full_image
        layout.addWidget(self.img_label, alignment=Qt.AlignCenter)

        # Buttons
        btn_layout = QHBoxLayout()
        btn_layout.setSpacing(2)

        rotate_btn = QPushButton("↻")
        rotate_btn.setFixedSize(30, 25)
        rotate_btn.clicked.connect(self._rotate)
        btn_layout.addWidget(rotate_btn)

        delete_btn = QPushButton("✕")
        delete_btn.setFixedSize(30, 25)
        delete_btn.setStyleSheet("background-color: #e74c3c; color: white;")
        delete_btn.clicked.connect(self._delete)
        btn_layout.addWidget(delete_btn)

        layout.addLayout(btn_layout)

        # Filename
        name = os.path.basename(photo.path)
        display_name = name[:14] + "..." if len(name) > 14 else name
        name_label = QLabel(display_name)
        name_label.setFont(QFont("Arial", 8))
        name_label.setAlignment(Qt.AlignCenter)
        layout.addWidget(name_label)

        # Load image
        self._load_image()

    def _load_image(self):
        pixmap = self.photo.get_pixmap()
        if pixmap:
            scaled = pixmap.scaled(100, 100, Qt.KeepAspectRatio, Qt.SmoothTransformation)
            self.img_label.setPixmap(scaled)
        else:
            self.img_label.setText("Err")

    def _show_full_image(self, event):
        """Ouvre la photo en grand dans une modal"""
        dialog = ImageViewerDialog(self.photo, self.window())
        dialog.exec_()

    def _rotate(self):
        self.photo.rotate()
        self._load_image()
        self.on_rotate(self.index)

    def _delete(self):
        self.on_delete(self.index)


class ExportThread(QThread):
    """Thread pour l'export Word"""
    progress = pyqtSignal(int)
    finished = pyqtSignal(str)
    error = pyqtSignal(str)

    def __init__(self, photos, path, ppp):
        super().__init__()
        self.photos = photos
        self.path = path
        self.ppp = ppp

    def run(self):
        try:
            self._generate_word()
            self.finished.emit(self.path)
        except Exception as e:
            self.error.emit(str(e))

    def _generate_word(self):
        doc = Document()

        # Marges de page Word
        page_margin = 10  # mm
        for section in doc.sections:
            section.top_margin = Mm(page_margin)
            section.bottom_margin = Mm(page_margin)
            section.left_margin = Mm(page_margin)
            section.right_margin = Mm(page_margin)

        cols, rows = {4: (2, 2), 6: (2, 3), 9: (3, 3)}[self.ppp]

        # Zone disponible dans Word (A4 = 210x297mm moins marges)
        available_w_mm = 210 - (page_margin * 2)  # 190mm
        available_h_mm = 297 - (page_margin * 2)  # 277mm

        # REDUIRE a 85% pour garantir que ca tienne TOUJOURS
        safe_factor = 0.85
        page_w_mm = available_w_mm * safe_factor
        page_h_mm = available_h_mm * safe_factor

        # Marge uniforme entre photos (horizontale et verticale identiques)
        gap_mm = 2  # Marge en mm entre chaque photo

        # Calculer taille des cellules en mm
        cell_w_mm = (page_w_mm - gap_mm * (cols - 1)) / cols
        cell_h_mm = (page_h_mm - gap_mm * (rows - 1)) / rows

        # Conversion mm -> pixels (150 DPI pour eviter les gros fichiers)
        dpi = 150
        mm_to_px = dpi / 25.4

        cell_w_px = int(cell_w_mm * mm_to_px)
        cell_h_px = int(cell_h_mm * mm_to_px)
        gap_px = int(gap_mm * mm_to_px)

        # Taille de l'image composite en pixels
        composite_w_px = cols * cell_w_px + (cols - 1) * gap_px
        composite_h_px = rows * cell_h_px + (rows - 1) * gap_px

        total = len(self.photos)
        num_pages = math.ceil(total / self.ppp)

        for page_idx in range(num_pages):
            if page_idx > 0:
                doc.add_page_break()

            # Creer image composite pour cette page
            composite = Image.new('RGB', (composite_w_px, composite_h_px), (255, 255, 255))

            start = page_idx * self.ppp
            photos_on_page = min(self.ppp, total - start)

            for i in range(rows):
                for j in range(cols):
                    idx = start + i * cols + j
                    if idx >= total:
                        continue

                    photo = self.photos[idx]

                    # Position dans l'image composite
                    x = j * (cell_w_px + gap_px)
                    y = i * (cell_h_px + gap_px)

                    try:
                        with Image.open(photo.path) as img:
                            if photo.rotation:
                                img = img.rotate(-photo.rotation, expand=True)

                            if img.mode != 'RGB':
                                if img.mode in ('RGBA', 'LA', 'P'):
                                    bg = Image.new('RGB', img.size, (255, 255, 255))
                                    if img.mode == 'P':
                                        img = img.convert('RGBA')
                                    if 'A' in img.getbands():
                                        bg.paste(img, mask=img.split()[-1])
                                    else:
                                        bg.paste(img)
                                    img = bg
                                else:
                                    img = img.convert('RGB')

                            # Crop-to-fill: recadrer pour remplir exactement la cellule
                            img_w, img_h = img.size
                            img_ratio = img_w / img_h
                            cell_ratio = cell_w_px / cell_h_px

                            if img_ratio > cell_ratio:
                                # Image trop large, couper les cotes
                                new_w = int(img_h * cell_ratio)
                                left = (img_w - new_w) // 2
                                img = img.crop((left, 0, left + new_w, img_h))
                            else:
                                # Image trop haute, couper haut/bas
                                new_h = int(img_w / cell_ratio)
                                top = (img_h - new_h) // 2
                                img = img.crop((0, top, img_w, top + new_h))

                            # Redimensionner a la taille exacte de la cellule
                            resample = Image.LANCZOS if hasattr(Image, 'LANCZOS') else Image.ANTIALIAS
                            img_resized = img.resize((cell_w_px, cell_h_px), resample)

                            # Placer directement sans centrage (remplit toute la cellule)
                            composite.paste(img_resized, (x, y))

                    except Exception:
                        pass  # Ignorer les erreurs

                    self.progress.emit(int((idx + 1) / total * 100))

            # Sauvegarder l'image composite et l'inserer dans le document
            buf = io.BytesIO()
            composite.save(buf, format='JPEG', quality=95)
            buf.seek(0)

            # Inserer avec les dimensions reduites (85% de la zone disponible)
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.space_before = Mm(0)
            para.paragraph_format.space_after = Mm(0)
            para.add_run().add_picture(buf, width=Mm(page_w_mm), height=Mm(page_h_mm))

        doc.save(self.path)


class PhotoManagerApp(QMainWindow):
    """Application principale PyQt5"""

    def __init__(self):
        super().__init__()
        self.setWindowTitle("Photo Manager")
        self.setMinimumSize(900, 600)
        self.resize(1200, 800)

        self.photos: List[PhotoItem] = []
        self.current_page = 0
        self._cards: List[PhotoCard] = []

        self._build_ui()

    def _build_ui(self):
        central = QWidget()
        self.setCentralWidget(central)

        main_layout = QHBoxLayout(central)
        main_layout.setContentsMargins(10, 10, 10, 10)
        main_layout.setSpacing(10)

        # Left panel
        left_panel = QFrame()
        left_panel.setFixedWidth(240)
        left_panel.setFrameStyle(QFrame.StyledPanel)
        left_layout = QVBoxLayout(left_panel)
        left_layout.setSpacing(5)

        # Title
        title = QLabel("Photo Manager")
        title.setFont(QFont("Arial", 18, QFont.Bold))
        title.setAlignment(Qt.AlignCenter)
        left_layout.addWidget(title)

        subtitle = QLabel("Export Word")
        subtitle.setAlignment(Qt.AlignCenter)
        left_layout.addWidget(subtitle)

        left_layout.addSpacing(15)

        # Add section
        add_label = QLabel("Ajouter")
        add_label.setFont(QFont("Arial", 11, QFont.Bold))
        left_layout.addWidget(add_label)

        folder_btn = QPushButton("+ Dossier")
        folder_btn.clicked.connect(self._add_folder)
        left_layout.addWidget(folder_btn)

        files_btn = QPushButton("+ Fichiers")
        files_btn.clicked.connect(self._add_files)
        left_layout.addWidget(files_btn)

        self.count_label = QLabel("0 photos")
        self.count_label.setFont(QFont("Arial", 10, QFont.Bold))
        self.count_label.setAlignment(Qt.AlignCenter)
        left_layout.addWidget(self.count_label)

        left_layout.addSpacing(10)

        # Photos per page
        ppp_label = QLabel("Photos/page")
        ppp_label.setFont(QFont("Arial", 11, QFont.Bold))
        left_layout.addWidget(ppp_label)

        self.ppp_group = QButtonGroup(self)
        for val, text in [(4, "4 (2x2)"), (6, "6 (2x3)"), (9, "9 (3x3)")]:
            radio = QRadioButton(text)
            radio.setProperty("ppp", val)
            self.ppp_group.addButton(radio, val)
            left_layout.addWidget(radio)
            if val == 6:
                radio.setChecked(True)

        left_layout.addSpacing(15)

        # Export button
        export_btn = QPushButton("EXPORTER WORD")
        export_btn.setFont(QFont("Arial", 12, QFont.Bold))
        export_btn.setStyleSheet("background-color: #27ae60; color: white; padding: 10px;")
        export_btn.clicked.connect(self._export)
        left_layout.addWidget(export_btn)

        clear_btn = QPushButton("Tout effacer")
        clear_btn.setStyleSheet("background-color: #e74c3c; color: white;")
        clear_btn.clicked.connect(self._clear)
        left_layout.addWidget(clear_btn)

        left_layout.addStretch()

        formats_label = QLabel("JPG, PNG")
        formats_label.setStyleSheet("color: gray;")
        formats_label.setAlignment(Qt.AlignCenter)
        left_layout.addWidget(formats_label)

        main_layout.addWidget(left_panel)

        # Right panel
        right_panel = QFrame()
        right_panel.setFrameStyle(QFrame.StyledPanel)
        right_layout = QVBoxLayout(right_panel)

        # Header with navigation
        header = QHBoxLayout()

        preview_label = QLabel("Apercu")
        preview_label.setFont(QFont("Arial", 13, QFont.Bold))
        header.addWidget(preview_label)

        header.addStretch()

        self.prev_btn = QPushButton("<")
        self.prev_btn.setFixedWidth(30)
        self.prev_btn.clicked.connect(self._prev_page)
        header.addWidget(self.prev_btn)

        self.page_label = QLabel("0/0")
        self.page_label.setFixedWidth(60)
        self.page_label.setAlignment(Qt.AlignCenter)
        header.addWidget(self.page_label)

        self.next_btn = QPushButton(">")
        self.next_btn.setFixedWidth(30)
        self.next_btn.clicked.connect(self._next_page)
        header.addWidget(self.next_btn)

        right_layout.addLayout(header)

        # Scroll area for photos
        self.scroll_area = QScrollArea()
        self.scroll_area.setWidgetResizable(True)
        self.scroll_area.setStyleSheet("background-color: #f5f5f5;")

        self.grid_widget = QWidget()
        self.grid_layout = QGridLayout(self.grid_widget)
        self.grid_layout.setSpacing(5)
        self.grid_layout.setAlignment(Qt.AlignTop | Qt.AlignLeft)

        self.scroll_area.setWidget(self.grid_widget)
        right_layout.addWidget(self.scroll_area)

        main_layout.addWidget(right_panel, 1)

    def _add_folder(self):
        folder = QFileDialog.getExistingDirectory(self, "Dossier")
        if folder:
            files = sorted([
                os.path.join(folder, f) for f in os.listdir(folder)
                if f.lower().endswith(SUPPORTED_FORMATS) and not f.startswith('.')
            ])
            if files:
                self._add_photos(files)

    def _add_files(self):
        files, _ = QFileDialog.getOpenFileNames(
            self, "Photos", "",
            "Images (*.jpg *.jpeg *.png *.JPG *.JPEG *.PNG)"
        )
        if files:
            self._add_photos(files)

    def _add_photos(self, files: List[str]):
        existing = {p.path for p in self.photos}
        new_photos = [PhotoItem(f) for f in files if f not in existing]
        self.photos.extend(new_photos)
        self._update()

    def _delete_photo(self, index: int):
        if 0 <= index < len(self.photos):
            self.photos[index].clear()
            del self.photos[index]
            max_page = max(0, (len(self.photos) - 1) // PHOTOS_PER_VIEW)
            if self.current_page > max_page:
                self.current_page = max_page
            self._update()

    def _rotate_photo(self, index: int):
        pass  # Rotation already handled in PhotoCard

    def _clear(self):
        for p in self.photos:
            p.clear()
        self.photos.clear()
        self.current_page = 0
        self._update()

    def _prev_page(self):
        if self.current_page > 0:
            self.current_page -= 1
            self._refresh_grid()

    def _next_page(self):
        max_page = max(0, (len(self.photos) - 1) // PHOTOS_PER_VIEW)
        if self.current_page < max_page:
            self.current_page += 1
            self._refresh_grid()

    def _update(self):
        self.count_label.setText(f"{len(self.photos)} photos")
        self._refresh_grid()

    def _refresh_grid(self):
        # Clear existing cards
        for card in self._cards:
            card.deleteLater()
        self._cards.clear()

        # Clear layout
        while self.grid_layout.count():
            item = self.grid_layout.takeAt(0)
            if item.widget():
                item.widget().deleteLater()

        if not self.photos:
            self.page_label.setText("0/0")
            return

        total_pages = math.ceil(len(self.photos) / PHOTOS_PER_VIEW)
        self.page_label.setText(f"{self.current_page + 1}/{total_pages}")

        start = self.current_page * PHOTOS_PER_VIEW
        end = min(start + PHOTOS_PER_VIEW, len(self.photos))

        cols = 6

        for i, idx in enumerate(range(start, end)):
            card = PhotoCard(
                self.photos[idx], idx,
                self._delete_photo, self._rotate_photo
            )
            self.grid_layout.addWidget(card, i // cols, i % cols)
            self._cards.append(card)

    def _export(self):
        if not self.photos:
            QMessageBox.warning(self, "Attention", "Aucune photo.")
            return

        path, _ = QFileDialog.getSaveFileName(
            self, "Enregistrer", "photos.docx",
            "Word (*.docx)"
        )
        if not path:
            return

        ppp = self.ppp_group.checkedId()

        progress = QProgressDialog("Generation du document Word...", None, 0, 100, self)
        progress.setWindowModality(Qt.WindowModal)
        progress.setAutoClose(True)
        progress.show()

        self.export_thread = ExportThread(self.photos, path, ppp)
        self.export_thread.progress.connect(progress.setValue)
        self.export_thread.finished.connect(lambda p: self._export_done(p, progress))
        self.export_thread.error.connect(lambda e: self._export_error(e, progress))
        self.export_thread.start()

    def _export_done(self, path, progress):
        progress.close()
        QMessageBox.information(self, "Succes", f"Exporte:\n{path}")

    def _export_error(self, error, progress):
        progress.close()
        QMessageBox.critical(self, "Erreur", error)


def main():
    app = QApplication(sys.argv)
    app.setStyle('Fusion')
    window = PhotoManagerApp()
    window.show()
    sys.exit(app.exec_())


if __name__ == "__main__":
    main()
